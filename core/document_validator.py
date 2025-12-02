"""
Strict Document Validator: Validates that uploaded documents are financial/deal documents (PDF or DOCX).

This module ensures only relevant financial deal documents are processed,
rejecting irrelevant files like blank documents, images, or unrelated content.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any, Tuple

from core.gemini_client import generate_json


# Financial/Deal keywords that MUST appear in valid documents
FINANCIAL_KEYWORDS = [
    # Deal Terms
    "agreement", "valuation", "acquirer", "consideration", "deal size",
    "company", "share", "financials", "investment", "term sheet",
    "equity", "debt", "purchase", "contract", "merger", "acquisition",
    
    # Financial Terms
    "revenue", "ebitda", "ebit", "profit", "margin", "cash flow",
    "assets", "liabilities", "balance sheet", "income statement",
    "valuation", "multiple", "discount", "premium", "synergy",
    
    # Deal Process
    "due diligence", "closing", "transaction", "buyer", "seller",
    "stakeholder", "shareholder", "board", "approval", "execution",
    
    # Sales Deal Terms (for lost deals)
    "deal", "opportunity", "proposal", "quote", "pricing", "negotiation",
    "customer", "client", "lost", "closed", "won", "rejected",
    "competitor", "alternative", "vendor", "solution",
    
    # Additional business context
    "business", "enterprise", "corporation", "partnership", "stake",
    "acquisition", "merger", "divestiture", "ipo", "private equity",
    "venture capital", "investment", "portfolio", "valuation model",
]

# Minimum required financial keywords for validation
MIN_FINANCIAL_KEYWORDS = 3

# Keywords that indicate NON-business/illegal content (should be minimal or absent)
NON_BUSINESS_KEYWORDS = [
    # Food & Cooking
    "recipe", "cooking", "food", "restaurant", "menu", "ingredient", "cuisine",
    "baking", "chef", "kitchen", "dining", "meal", "dish", "flavor",
    
    # Fiction & Literature
    "novel", "fiction", "story", "chapter", "character", "plot", "narrative",
    "poem", "poetry", "verse", "stanza", "prose", "literature", "author",
    
    # Medical & Health
    "medical", "prescription", "diagnosis", "treatment", "patient", "doctor",
    "hospital", "clinic", "symptom", "disease", "medication", "therapy",
    "surgery", "physician", "nurse", "healthcare",
    
    # Academic & Research
    "academic", "research paper", "thesis", "dissertation", "essay", "citation",
    "bibliography", "abstract", "methodology", "hypothesis", "experiment",
    "university", "college", "scholar", "professor", "student",
    
    # Legal (non-business)
    "legal brief", "court", "lawsuit", "litigation", "judgment", "verdict",
    "plaintiff", "defendant", "attorney", "lawyer", "legal case",
    
    # Personal & Social
    "resume", "cv", "curriculum vitae", "job application", "cover letter",
    "personal", "family", "relationship", "dating", "marriage",
    
    # Entertainment
    "movie", "film", "cinema", "actor", "actress", "director", "script",
    "music", "song", "album", "artist", "concert", "performance",
    
    # Sports & Hobbies (only if clearly non-business)
    "sports", "game", "player", "match", "tournament", "championship",
    "hobby", "leisure", "recreation", "entertainment",
    
    # News & Media (only if clearly non-business)
    "news", "article", "blog", "journalism", "reporter", "journalist",
    "newspaper", "magazine", "editorial", "headline",
    
    # Technical Manuals (non-business)
    "manual", "tutorial", "how-to", "guide", "instruction", "troubleshooting",
    "user guide", "technical documentation", "api documentation",
    
    # Certificates & Education
    "certificate", "diploma", "degree", "graduation", "education", "school",
    "course", "curriculum", "syllabus",
    
    # Forms & Applications (only if clearly non-business)
    "survey", "questionnaire", "poll", "feedback",
    
    # Illegal/Inappropriate Content Indicators
    "illegal", "criminal", "fraud", "scam", "piracy", "copyright infringement",
    "hack", "malware", "virus", "exploit", "unauthorized access",
]

# Maximum allowed non-business keywords (if exceeded, likely not a business doc)
MAX_NON_BUSINESS_KEYWORDS = 5  # More lenient - allow up to 5 non-business keywords

# Minimum text length
MIN_TEXT_LENGTH = 200

# Maximum file size (100 MB)
MAX_FILE_SIZE = 100 * 1024 * 1024

# Minimum file size (100 bytes)
MIN_FILE_SIZE = 100


def validate_file_type(filename: str) -> Tuple[bool, str]:
    """
    Validate that the file is a PDF or DOCX.
    
    Args:
        filename: Name of the uploaded file
        
    Returns:
        Tuple of (is_valid: bool, error_message: str)
    """
    file_path = Path(filename)
    extension = file_path.suffix.lower()
    
    allowed_extensions = {".pdf", ".docx", ".doc", ".txt"}
    
    if extension not in allowed_extensions:
        return False, (
            f"❌ Invalid File Type – Only PDF, DOCX, and TXT files are allowed.\n\n"
            f"Uploaded file type: {extension or 'unknown'}\n"
            f"Please upload a PDF, DOCX, or TXT file (e.g., valuation report, term sheet, M&A document, financial contract)."
        )
    
    return True, ""


def extract_text_from_document(file_bytes: bytes, filename: str) -> Tuple[str, int]:
    """
    Extract text from PDF, DOCX, or TXT and count pages/sections.
    
    Args:
        file_bytes: Document file content as bytes
        filename: Name of the file (for extension detection)
        
    Returns:
        Tuple of (extracted_text: str, page_count: int)
    """
    file_path = Path(filename)
    extension = file_path.suffix.lower()
    
    if extension == ".pdf":
        return _extract_text_from_pdf(file_bytes)
    elif extension in {".docx", ".doc"}:
        return _extract_text_from_docx(file_bytes)
    elif extension == ".txt":
        return _extract_text_from_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {extension}")


def _extract_text_from_pdf(file_bytes: bytes) -> Tuple[str, int]:
    """
    Extract text from PDF and count pages.
    
    Args:
        file_bytes: PDF file content as bytes
        
    Returns:
        Tuple of (extracted_text: str, page_count: int)
    """
    try:
        from pypdf import PdfReader
        import io
        
        pdf_file = io.BytesIO(file_bytes)
        reader = PdfReader(pdf_file)
        
        page_count = len(reader.pages)
        text_parts = []
        
        for page_num, page in enumerate(reader.pages, 1):
            try:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            except Exception:
                # Skip pages that can't be extracted
                continue
        
        extracted_text = "\n\n".join(text_parts)
        return extracted_text, page_count
        
    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")


def _extract_text_from_docx(file_bytes: bytes) -> Tuple[str, int]:
    """
    Extract text from DOCX and estimate page count.
    
    Args:
        file_bytes: DOCX file content as bytes
        
    Returns:
        Tuple of (extracted_text: str, estimated_page_count: int)
    """
    try:
        # Use the same loader as the main system for consistency
        from langchain_community.document_loaders import Docx2txtLoader
        import tempfile
        import os
        
        # Save to temp file for Docx2txtLoader
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
            temp_file.write(file_bytes)
            temp_path = temp_file.name
        
        try:
            loader = Docx2txtLoader(temp_path)
            documents = loader.load()
            
            # Combine all document text
            text_parts = [doc.page_content for doc in documents if doc.page_content.strip()]
            extracted_text = "\n\n".join(text_parts)
            
            # Estimate page count (approximately 500 words per page)
            word_count = len(extracted_text.split())
            estimated_pages = max(1, word_count // 500)
            
            return extracted_text, estimated_pages
        finally:
            # Clean up temp file
            if os.path.exists(temp_path):
                os.unlink(temp_path)
        
    except Exception as e:
        # Fallback: try python-docx if available
        try:
            from docx import Document
            import io
            
            docx_file = io.BytesIO(file_bytes)
            doc = Document(docx_file)
            
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            
            extracted_text = "\n\n".join(text_parts)
            
            # Estimate page count
            word_count = len(extracted_text.split())
            estimated_pages = max(1, word_count // 500)
            
            return extracted_text, estimated_pages
        except Exception as e2:
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}. Alternative method also failed: {str(e2)}")


def _extract_text_from_txt(file_bytes: bytes) -> Tuple[str, int]:
    """
    Extract text from TXT file.
    
    Args:
        file_bytes: TXT file content as bytes
        
    Returns:
        Tuple of (extracted_text: str, estimated_page_count: int)
    """
    try:
        # Decode text from bytes
        extracted_text = file_bytes.decode('utf-8', errors='ignore')
        
        # Estimate page count (approximately 500 words per page)
        word_count = len(extracted_text.split())
        estimated_pages = max(1, word_count // 500)
        
        return extracted_text, estimated_pages
    except Exception as e:
        raise ValueError(f"Failed to extract text from TXT: {str(e)}")


def validate_financial_document(text: str, page_count: int = 1) -> Tuple[bool, str]:
    """
    Validate that the document contains financial/deal content.
    
    Args:
        text: Extracted text from document
        page_count: Number of pages/sections in document
        
    Returns:
        Tuple of (is_valid: bool, error_message: str)
    """
    if not text or len(text.strip()) < MIN_TEXT_LENGTH:
        return False, (
            f"❌ Invalid Document – Document text is too short.\n\n"
            f"Found: {len(text.strip())} characters (minimum {MIN_TEXT_LENGTH} required).\n"
            f"This document may be blank, image-based, or corrupted.\n"
            f"Please upload a document with extractable text content."
        )
    
    text_lower = text.lower()
    
    # Check for financial keywords
    found_keywords = [kw for kw in FINANCIAL_KEYWORDS if kw in text_lower]
    keyword_count = len(found_keywords)
    
    if keyword_count < MIN_FINANCIAL_KEYWORDS:
        return False, (
            f"❌ Invalid Document – This PDF does not appear to be a deal/financial document.\n\n"
            f"Found only {keyword_count} financial/deal-related terms (minimum {MIN_FINANCIAL_KEYWORDS} required).\n"
            f"Please upload a contract, valuation report, term sheet, or M&A document.\n\n"
            f"Expected keywords include: agreement, valuation, deal, contract, financials, investment, etc."
        )
    
    # Check for currency symbols or numbers (financial indicators)
    currency_patterns = [r'\$', r'₹', r'€', r'£', r'¥', r'%', r'\d+[,\d]*\.?\d*\s*(?:million|billion|thousand|k|m|b)', r'\d+[,\d]*\.?\d*']
    has_currency = any(re.search(pattern, text, re.IGNORECASE) for pattern in currency_patterns)
    
    if not has_currency:
        return False, (
            f"❌ Invalid Document – Document does not contain financial indicators.\n\n"
            f"No currency symbols ($, ₹, €, £) or financial numbers found.\n"
            f"Please upload a financial document with monetary values or percentages."
        )
    
    # Check for structured text (at least 1-2 lines with substantial content)
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    substantial_lines = [line for line in lines if len(line) > 15]  # Lines with >15 chars (reduced from 20)
    
    if len(substantial_lines) < 1:
        return False, (
            f"❌ Invalid Document – Document lacks structured content.\n\n"
            f"Found only {len(substantial_lines)} substantial lines (minimum 1 required).\n"
            f"This document may be mostly blank or contain only images.\n"
            f"Please upload a document with readable, structured text content."
        )
    
    # Check page count (should have at least 1 page/section)
    if page_count < 1:
        return False, (
            f"❌ Invalid Document – Document has no readable content.\n\n"
            f"Please upload a valid document file with extractable content."
        )
    
    # Lenient validation: Check for non-business/illegal content indicators
    # Only flag if there are MANY non-business keywords (likely not a business doc)
    non_business_count = sum(1 for keyword in NON_BUSINESS_KEYWORDS if keyword in text_lower)
    
    # Only reject if there are significantly more non-business keywords than business keywords
    financial_keyword_count = sum(1 for keyword in FINANCIAL_KEYWORDS if keyword in text_lower)
    
    # Reject only if non-business keywords significantly outweigh business keywords
    if non_business_count > MAX_NON_BUSINESS_KEYWORDS and non_business_count > financial_keyword_count * 2:
        # Identify which non-business keywords were found
        found_keywords = [kw for kw in NON_BUSINESS_KEYWORDS if kw in text_lower]
        keyword_examples = ", ".join(found_keywords[:5])  # Show first 5
        
        return False, (
            f"❌ Invalid Document – Document contains excessive non-business content.\n\n"
            f"Found {non_business_count} non-business content indicators vs {financial_keyword_count} business terms.\n"
            f"Detected keywords: {keyword_examples}{'...' if len(found_keywords) > 5 else ''}\n\n"
            f"This system ONLY processes financial/deal documents.\n"
            f"Please upload a contract, valuation report, term sheet, or M&A document."
        )
    
    # Additional check: Ensure business context is present
    business_context_keywords = [
        "business", "company", "corporation", "enterprise", "organization",
        "deal", "transaction", "agreement", "contract", "financial",
        "investment", "acquisition", "merger", "valuation", "revenue"
    ]
    business_context_count = sum(1 for kw in business_context_keywords if kw in text_lower)
    
    if business_context_count < 2:
        return False, (
            f"❌ Invalid Document – Document lacks business context.\n\n"
            f"Found only {business_context_count} business-related terms (minimum 2 required).\n"
            f"This document does not appear to be a business/financial document.\n"
            f"Please upload a contract, valuation report, term sheet, or M&A document."
        )
    
    return True, ""


def validate_file_size(file_bytes: bytes) -> Tuple[bool, str]:
    """
    Validate file size constraints.
    
    Args:
        file_bytes: File content as bytes
        
    Returns:
        Tuple of (is_valid: bool, error_message: str)
    """
    file_size = len(file_bytes)
    
    if file_size < MIN_FILE_SIZE:
        return False, (
            f"❌ Invalid File – File is too small or empty.\n\n"
            f"File size: {file_size} bytes (minimum {MIN_FILE_SIZE} bytes required)."
        )
    
    if file_size > MAX_FILE_SIZE:
        size_mb = file_size / (1024 * 1024)
        max_mb = MAX_FILE_SIZE / (1024 * 1024)
        return False, (
            f"❌ Invalid File – File is too large.\n\n"
            f"File size: {size_mb:.1f} MB (maximum {max_mb} MB allowed).\n"
            f"Please upload a smaller file or compress the document."
        )
    
    return True, ""


def sanitize_pdf_text(text: str) -> str:
    """
    Sanitize PDF text by removing excessive whitespace and cleaning up.
    
    Args:
        text: Raw extracted text
        
    Returns:
        Sanitized text
    """
    # Remove excessive whitespace
    text = re.sub(r'\s+', ' ', text)
    # Remove control characters except newlines
    text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F]', '', text)
    # Normalize line breaks
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def validate_document_relevance(text: str, filename: str = "") -> Tuple[bool, str]:
    """
    Final LLM-based validation for edge cases with strict business content checking.
    
    Args:
        text: Extracted text from document
        filename: Optional filename
        
    Returns:
        Tuple of (is_valid: bool, error_message: str)
    """
    if not text or len(text.strip()) < MIN_TEXT_LENGTH:
        return False, "Document is too short or empty."
    
    text_lower = text.lower()
    
    # STRICT keyword check first
    keyword_count = sum(1 for keyword in FINANCIAL_KEYWORDS if keyword in text_lower)
    if keyword_count < MIN_FINANCIAL_KEYWORDS:
        return False, (
            f"Document does not contain sufficient financial/deal-related content. "
            f"Found {keyword_count} relevant terms (minimum {MIN_FINANCIAL_KEYWORDS} required)."
        )
    
    # Check for non-business content (lenient - only reject if significantly more non-business than business)
    non_business_count = sum(1 for keyword in NON_BUSINESS_KEYWORDS if keyword in text_lower)
    financial_keyword_count = sum(1 for keyword in FINANCIAL_KEYWORDS if keyword in text_lower)
    
    # Only reject if non-business keywords significantly outweigh business keywords
    if non_business_count > MAX_NON_BUSINESS_KEYWORDS and non_business_count > financial_keyword_count * 2:
        found_keywords = [kw for kw in NON_BUSINESS_KEYWORDS if kw in text_lower]
        return False, (
            f"Document contains excessive non-business content. "
            f"Found {non_business_count} non-business indicators vs {financial_keyword_count} business terms: {', '.join(found_keywords[:3])}... "
        )
    
    # LLM validation for edge cases (STRICT mode)
    validation_result = _llm_validate_document(text[:4000], filename)  # Increased context
    
    if not validation_result.get("is_deal_document", False):
        reason = validation_result.get("reason", "Document does not appear to be a financial/deal document.")
        return False, reason
    
    # Additional confidence check (more lenient)
    confidence = validation_result.get("confidence", 0.5)
    if confidence < 0.4:  # Require at least 40% confidence (reduced from 70%)
        return False, (
            f"Document validation confidence too low ({confidence:.1%}). "
            f"Please ensure the document is a financial/deal document."
        )
    
    return True, ""


def _llm_validate_document(text: str, filename: str = "") -> dict[str, Any]:
    """
    Use LLM to validate if document is financial/deal-related with lenient validation.
    
    Args:
        text: Document text (first 4000 chars)
        filename: Optional filename
        
    Returns:
        Dictionary with validation result
    """
    prompt = f"""
You are a document validator for a Deal Forensics AI system that analyzes financial deals and M&A transactions.

Your task: Determine if the following document is a financial/deal document. Be reasonable and lenient.

CRITICAL REQUIREMENTS - This system ONLY processes:
- Financial contracts and agreements (with deal terms, pricing, valuation)
- Valuation reports and term sheets (with financial metrics)
- M&A documents and acquisition agreements (with transaction details)
- Investment proposals and deal summaries (with investment terms)
- Lost sales deal post-mortems (with financial context, pricing, competitors)
- Deal failure analysis with financial terms (revenue, pricing, budget)

CRITICAL REJECTIONS - This system does NOT process:
- Random PDFs, invoices, receipts (unless clearly deal-related with financial terms)
- Academic papers, research documents, theses, dissertations
- Novels, stories, fiction, poetry, literature
- Recipes, cooking guides, food blogs, restaurant menus
- Medical records, prescriptions, health documents
- Personal documents (resumes, CVs, personal letters)
- News articles, blog posts, journalism (unless about specific deals)
- Technical manuals, tutorials, how-to guides
- Entertainment content (movies, music, sports)
- Legal documents unrelated to business deals
- Forms, surveys, questionnaires
- Certificates, diplomas, educational documents
- ANY content without clear business/financial/deal context

VALIDATION RULES:
1. Document should contain financial terms (revenue, valuation, pricing, investment, etc.) OR business context (company, deal, transaction, agreement)
2. If document mentions deals, contracts, pricing, customers, revenue, or business transactions, ACCEPT it
3. Only REJECT if document is clearly about recipes, fiction, medical records, academic papers, or completely unrelated topics
4. If uncertain, ACCEPT the document (be lenient)

Document Content (first 4000 characters):
{text[:4000]}

Filename: {filename or "Unknown"}

Analyze the document and return JSON:
{{
  "is_deal_document": true/false,
  "reason": "Brief explanation of why it is or isn't a financial/deal document.",
  "confidence": 0.0-1.0
}}

Be LENIENT - return true if this appears to be a financial/deal document with business/financial content, even if it's not perfect.
Only return false if this is CLEARLY about recipes, fiction, medical records, academic papers, personal documents, or completely unrelated topics.
If the document mentions deals, contracts, pricing, customers, revenue, business transactions, or sales, it should be ACCEPTED.
If uncertain, default to ACCEPTING the document.
""".strip()
    
    try:
        result = generate_json(prompt)
        if not isinstance(result, dict):
            # If LLM fails, default to accepting (lenient)
            return {"is_deal_document": True, "reason": "LLM validation unavailable, defaulting to accept", "confidence": 0.6}
        
        # If LLM says false but confidence is low, be lenient and accept
        if not result.get("is_deal_document", False) and result.get("confidence", 0.5) < 0.6:
            return {"is_deal_document": True, "reason": "Low confidence rejection, defaulting to accept", "confidence": 0.5}
        
        return result
    except Exception:
        # If LLM fails, default to accepting (lenient approach)
        return {"is_deal_document": True, "reason": "LLM validation unavailable, defaulting to accept", "confidence": 0.6}


def check_document_quality(text: str) -> Tuple[bool, str]:
    """
    Check document quality and provide warnings with strict business content validation.
    
    Args:
        text: Document text
        
    Returns:
        Tuple of (is_acceptable: bool, warning_message: str)
    """
    if not text:
        return False, "Document is empty."
    
    if len(text) < MIN_TEXT_LENGTH:
        return False, (
            f"Document is too short. Financial/deal documents should be at least {MIN_TEXT_LENGTH} characters "
            f"for meaningful analysis. Found: {len(text)} characters."
        )
    
    # Check for excessive whitespace (likely corrupted or image-based PDF)
    whitespace_ratio = len(re.findall(r'\s+', text)) / len(text) if text else 0
    if whitespace_ratio > 0.5:
        return False, (
            "Document appears to be corrupted or image-based PDF with minimal extractable text. "
            "Please ensure the PDF contains selectable text, not just images."
        )
    
    text_lower = text.lower()
    
    # STRICT check for financial indicators
    financial_indicators = [
        "pricing", "valuation", "deal", "contract", "agreement", "financial",
        "investment", "acquisition", "merger", "equity", "revenue", "transaction",
        "business", "company", "corporation", "enterprise"
    ]
    found_indicators = sum(1 for indicator in financial_indicators if indicator in text_lower)
    
    if found_indicators < 1:  # Reduced to 1 for more lenient validation
        return False, (
            f"Document lacks sufficient financial/deal-related content. "
            f"Found only {found_indicators} business/financial indicators (minimum 1 required). "
            f"Please ensure the document discusses deals, contracts, valuations, or financial transactions."
        )
    
    # Check for non-business content (lenient - only reject if significantly more non-business than business)
    non_business_count = sum(1 for keyword in NON_BUSINESS_KEYWORDS if keyword in text_lower)
    financial_keyword_count = sum(1 for keyword in FINANCIAL_KEYWORDS if keyword in text_lower)
    
    # Only reject if non-business keywords significantly outweigh business keywords
    if non_business_count > MAX_NON_BUSINESS_KEYWORDS and non_business_count > financial_keyword_count * 2:
        found_keywords = [kw for kw in NON_BUSINESS_KEYWORDS if kw in text_lower]
        return False, (
            f"Document contains excessive non-business content. "
            f"Found {non_business_count} non-business indicators vs {financial_keyword_count} business terms: {', '.join(found_keywords[:3])}... "
            f"This system only processes financial/deal documents."
        )
    
    return True, ""


def handle_invalid_document(error_message: str) -> dict[str, Any]:
    """
    Handle invalid document by returning a structured error response.
    
    Args:
        error_message: Error message to return
        
    Returns:
        Dictionary with error information
    """
    return {
        "error": True,
        "error_message": error_message,
        "valid": False,
        "timeline": {},
        "comparative": {},
        "playbook": {},
        "scorecard": {},
        "metadata": {},
        "report": None,
    }


def process_valid_document(text: str, page_count: int) -> Tuple[bool, str]:
    """
    Final validation check before processing.
    
    Args:
        text: Extracted text
        page_count: Number of pages/sections
        
    Returns:
        Tuple of (is_valid: bool, error_message: str)
    """
    # Run all validation checks
    is_quality_ok, quality_msg = check_document_quality(text)
    if not is_quality_ok:
        return False, quality_msg
    
    is_financial, financial_msg = validate_financial_document(text, page_count)
    if not is_financial:
        return False, financial_msg
    
    return True, ""
